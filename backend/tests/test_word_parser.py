"""Tests for structure-preserving Word document parsing."""

import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.knowledge.parser import KnowledgeParser, WordParser


class WordParserTests(unittest.TestCase):
    """Verify DOCX parsing for headings, paragraphs, and tables."""

    def test_parses_word_document_sections_and_tables(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            topic_directory = root / "policies"
            topic_directory.mkdir()
            document_path = topic_directory / "onboarding.docx"

            word_document = Document()
            word_document.add_heading("Giriş", level=1)
            word_document.add_paragraph("Çalışan onboarding süreci açıklaması.")
            word_document.add_heading("Kontrol Listesi", level=2)
            word_document.add_paragraph("Belgeler işe başlamadan hazırlanır.")
            table = word_document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Adım"
            table.cell(0, 1).text = "Sorumlu"
            table.cell(1, 0).text = "Hesap açılışı"
            table.cell(1, 1).text = "BT"
            word_document.save(document_path)

            parsed_document = WordParser(root, language="tr").parse(
                document_path
            )

        self.assertEqual(parsed_document.document_name, "onboarding.docx")
        self.assertEqual(parsed_document.relative_path, "policies/onboarding.docx")
        self.assertEqual(parsed_document.language, "tr")
        self.assertIn("Çalışan onboarding süreci", parsed_document.content)
        self.assertIn("Adım | Sorumlu", parsed_document.content)
        self.assertEqual(
            [section.title for section in parsed_document.sections],
            ["Giriş", "Kontrol Listesi"],
        )
        self.assertEqual(parsed_document.sections[0].level, 1)
        self.assertEqual(parsed_document.sections[1].level, 2)
        self.assertIn(
            "Hesap açılışı | BT",
            parsed_document.sections[1].content,
        )

    def test_rejects_word_document_outside_knowledge_base(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            knowledge_base = root / "knowledge_base"
            knowledge_base.mkdir()
            outside_path = root / "outside.docx"
            document = Document()
            document.add_paragraph("Outside")
            document.save(outside_path)

            with self.assertRaises(ValueError):
                WordParser(knowledge_base).parse(outside_path)

    def test_rejects_invalid_docx_content(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            invalid_path = root / "invalid.docx"
            invalid_path.write_text("not a real docx", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "Unable to read Word"):
                WordParser(root).parse(invalid_path)

    def test_knowledge_parser_dispatches_docx_files(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            document_path = root / "handbook.docx"
            document = Document()
            document.add_heading("Handbook", level=1)
            document.add_paragraph("Company handbook content.")
            document.save(document_path)

            parsed_document = KnowledgeParser(root).parse(document_path)

        self.assertEqual(parsed_document.document_name, "handbook.docx")
        self.assertEqual(parsed_document.sections[0].title, "Handbook")


if __name__ == "__main__":
    unittest.main()
