"""Tests for full and incremental knowledge indexing."""

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.knowledge.cache import CorruptedCacheError, IndexCache
from app.knowledge.indexer import KnowledgeIndexer, KnowledgeIndexingError
from app.knowledge.loader import KnowledgeLoader
from app.knowledge.manifest import ManifestError, ManifestLoader

MANIFEST = """\
version: 1
default_language: en
chunk_size: 80
chunk_overlap: 10
supported_extensions:
  - docx
  - md
"""


def create_indexer(
    knowledge_base_path: Path,
    manifest_path: Path,
    cache_path: Path,
) -> KnowledgeIndexer:
    """Create a real indexer using temporary filesystem dependencies."""
    return KnowledgeIndexer(
        loader=KnowledgeLoader(knowledge_base_path),
        manifest_loader=ManifestLoader(manifest_path),
        cache=IndexCache(cache_path),
    )


class KnowledgeIndexerTests(unittest.TestCase):
    """Verify indexing orchestration, reports, and incremental behavior."""

    def setUp(self) -> None:
        self._temporary_directory = tempfile.TemporaryDirectory()
        self.root = Path(self._temporary_directory.name)
        self.knowledge_base = self.root / "knowledge_base"
        self.knowledge_base.mkdir()
        self.manifest_path = self.knowledge_base / "manifest.yaml"
        self.manifest_path.write_text(MANIFEST, encoding="utf-8")
        self.cache_path = self.root / "data" / "index_cache.json"

    def tearDown(self) -> None:
        self._temporary_directory.cleanup()

    def write_document(self, relative_path: str, content: str) -> Path:
        """Write one temporary knowledge text document."""
        path = self.knowledge_base / relative_path
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
        return path

    def write_word_document(
        self,
        relative_path: str,
        heading: str,
        content: str,
    ) -> Path:
        """Write one temporary Word knowledge document."""
        path = self.knowledge_base / relative_path
        path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading(heading, level=1)
        document.add_paragraph(content)
        document.save(path)
        return path

    def create_indexer(self) -> KnowledgeIndexer:
        """Create the test indexer."""
        return create_indexer(
            self.knowledge_base,
            self.manifest_path,
            self.cache_path,
        )

    def test_indexes_documents_and_generates_serializable_report(self) -> None:
        self.write_document(
            "python/variables.md",
            "# Variables\n\nVariables bind names to values.",
        )
        self.write_document(
            "fastapi/routing.md",
            "# Routing\n\nRoutes connect paths to handlers.",
        )

        result = self.create_indexer().index()

        self.assertEqual(result.statistics.files_scanned, 2)
        self.assertEqual(result.statistics.files_indexed, 2)
        self.assertEqual(result.statistics.files_skipped, 0)
        self.assertEqual(result.statistics.files_removed, 0)
        self.assertEqual(
            result.statistics.chunks_created,
            len(result.chunks),
        )
        self.assertIn("Knowledge Index Report", result.statistics.to_report())
        serialized = json.loads(result.model_dump_json())
        self.assertEqual(serialized["statistics"]["files_indexed"], 2)

    def test_indexes_markdown_and_word_documents(self) -> None:
        self.write_document(
            "python/variables.md",
            "# Variables\n\nVariables bind names to values.",
        )
        self.write_word_document(
            "policies/onboarding.docx",
            "Onboarding",
            "Employees receive access before the first day.",
        )

        result = self.create_indexer().index()

        indexed_paths = [
            indexed_file.relative_path
            for indexed_file in result.indexed_files
        ]
        self.assertEqual(
            indexed_paths,
            ["policies/onboarding.docx", "python/variables.md"],
        )
        self.assertIn(
            "Employees receive access",
            "\n".join(chunk.content for chunk in result.chunks),
        )

    def test_skips_unchanged_documents(self) -> None:
        self.write_document("python/oop.md", "# OOP\n\nObjects and classes.")
        indexer = self.create_indexer()
        indexer.index()

        result = indexer.index()

        self.assertEqual(result.statistics.files_scanned, 1)
        self.assertEqual(result.statistics.files_indexed, 0)
        self.assertEqual(result.statistics.files_skipped, 1)
        self.assertEqual(result.statistics.chunks_created, 0)
        self.assertEqual(result.chunks, [])

    def test_reindexes_only_changed_documents(self) -> None:
        changed_path = self.write_document(
            "python/oop.md",
            "# OOP\n\nOriginal content.",
        )
        self.write_document(
            "git/commits.md",
            "# Commits\n\nSnapshots of changes.",
        )
        indexer = self.create_indexer()
        indexer.index()
        changed_path.write_text(
            "# OOP\n\nUpdated content.",
            encoding="utf-8",
        )

        result = indexer.index()

        self.assertEqual(result.statistics.files_indexed, 1)
        self.assertEqual(result.statistics.files_skipped, 1)
        self.assertEqual(
            [item.relative_path for item in result.indexed_files],
            ["python/oop.md"],
        )

    def test_detects_and_removes_deleted_documents(self) -> None:
        removed_path = self.write_document(
            "python/oop.md",
            "# OOP\n\nObjects and classes.",
        )
        self.write_document(
            "git/commits.md",
            "# Commits\n\nSnapshots of changes.",
        )
        indexer = self.create_indexer()
        indexer.index()
        removed_path.unlink()

        result = indexer.index()
        cached_state = IndexCache(self.cache_path).load()

        self.assertEqual(result.removed_files, ["python/oop.md"])
        self.assertEqual(result.statistics.files_removed, 1)
        self.assertNotIn("python/oop.md", cached_state.files)
        self.assertIn("git/commits.md", cached_state.files)

    def test_manifest_change_forces_safe_reindexing(self) -> None:
        self.write_document(
            "python/oop.md",
            "# OOP\n\nObjects and classes.",
        )
        indexer = self.create_indexer()
        indexer.index()
        self.manifest_path.write_text(
            MANIFEST.replace("chunk_size: 80", "chunk_size: 100"),
            encoding="utf-8",
        )

        result = indexer.index()

        self.assertEqual(result.statistics.files_indexed, 1)
        self.assertEqual(result.statistics.files_skipped, 0)

    def test_uses_manifest_chunk_configuration(self) -> None:
        self.write_document(
            "python/long_document.md",
            "# Long Document\n\n" + "A" * 300,
        )
        indexer = self.create_indexer()

        small_chunk_result = indexer.index()
        self.manifest_path.write_text(
            MANIFEST.replace("chunk_size: 80", "chunk_size: 500"),
            encoding="utf-8",
        )
        large_chunk_result = indexer.index()

        self.assertGreater(
            small_chunk_result.statistics.chunks_created,
            large_chunk_result.statistics.chunks_created,
        )
        self.assertEqual(
            large_chunk_result.statistics.chunks_created,
            1,
        )

    def test_reports_unreadable_markdown(self) -> None:
        path = self.knowledge_base / "python" / "invalid.md"
        path.parent.mkdir()
        path.write_bytes(b"\xff\xfe")

        with self.assertRaisesRegex(
            KnowledgeIndexingError,
            "python/invalid.md",
        ):
            self.create_indexer().index()

    def test_reports_missing_knowledge_directory(self) -> None:
        missing_path = self.root / "missing"
        indexer = create_indexer(
            missing_path,
            self.manifest_path,
            self.cache_path,
        )

        with self.assertRaisesRegex(
            FileNotFoundError,
            "Knowledge base directory not found",
        ):
            indexer.index()

    def test_propagates_invalid_manifest_error(self) -> None:
        self.manifest_path.write_text("version: [", encoding="utf-8")

        with self.assertRaises(ManifestError):
            self.create_indexer().index()

    def test_propagates_corrupted_cache_error(self) -> None:
        self.cache_path.parent.mkdir()
        self.cache_path.write_text("{invalid", encoding="utf-8")

        with self.assertRaises(CorruptedCacheError):
            self.create_indexer().index()


if __name__ == "__main__":
    unittest.main()
